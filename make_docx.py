"""
make_docx.py — converts the Substack article markdown to a polished Word doc.
Run from the project root:
  .venv/bin/python3 make_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "/Users/ruoqianzhu/Documents/heat_defense_substack.docx"

# ── colour palette ──────────────────────────────────────────────────────────
HEAT_RED   = RGBColor(0x98, 0x00, 0x2E)   # Heat dark red
HEAT_BLACK = RGBColor(0x14, 0x14, 0x14)
RULE_GREY  = RGBColor(0xCC, 0xCC, 0xCC)
TABLE_HEADER_BG = "981F2A"                 # hex string for shading (no #)
TABLE_ROW_ALT   = "F7F2F2"
PULL_QUOTE_BG   = "FFF5F5"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right  →  dict(sz, color, val)"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in kwargs.items():
        el = OxmlElement(f"w:{ side}")
        el.set(qn("w:val"),   attrs.get("val",   "single"))
        el.set(qn("w:sz"),    str(attrs.get("sz", 4)))
        el.set(qn("w:color"), attrs.get("color", "auto"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc, color_hex="CCCCCC"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pb.append(bottom)
    pPr.append(pb)
    return p


def fmt_run(run, bold=False, italic=False, color=None, size=None, font_name=None):
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name


def heading(doc, text, level=1, color=HEAT_RED, size=None, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12)
    run.font.name = "Georgia"
    return p


def subheading(doc, text):
    return heading(doc, text, level=2, color=HEAT_RED, size=13, space_before=14, space_after=4)


def body(doc, text, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    # parse inline bold/italic markers
    _add_inline(p, text, italic_default=italic)
    for run in p.runs:
        if not run.font.name:
            run.font.name = "Calibri"
        if not run.font.size:
            run.font.size = Pt(11)
        if not run.font.color.type:
            run.font.color.rgb = HEAT_BLACK
    return p


def _add_inline(p, text: str, italic_default=False):
    """Parse **bold** and *italic* markers and add runs."""
    import re
    # pattern: **bold** or *italic*
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*')
    pos = 0
    for m in pattern.finditer(text):
        # text before match
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            r.italic = italic_default
        if m.group(1) is not None:   # bold
            r = p.add_run(m.group(1))
            r.bold = True
            r.italic = italic_default
        else:                         # italic
            r = p.add_run(m.group(2))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        r.italic = italic_default


def pull_quote(doc, text):
    """Shaded callout box with left border — for the blockquote passages."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.4)
    p.paragraph_format.right_indent  = Inches(0.4)
    p.paragraph_format.space_before  = Pt(10)
    p.paragraph_format.space_after   = Pt(10)
    # left border via paragraph border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "981F2A")
    pb.append(left)
    pPr.append(pb)
    _add_inline(p, text, italic_default=True)
    for run in p.runs:
        run.font.name  = "Georgia"
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x50, 0x10, 0x10)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    _add_inline(p, text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = HEAT_BLACK
    return p


def data_table(doc, headers, rows, col_widths=None):
    """Styled data table with Heat-red header row and alternating row shading."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], TABLE_HEADER_BG)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold  = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].paragraph_format.space_before = Pt(3)
        hdr_cells[i].paragraphs[0].paragraph_format.space_after  = Pt(3)

    # data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg    = TABLE_ROW_ALT if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            cells[c_idx].paragraphs[0].paragraph_format.space_before = Pt(2)
            cells[c_idx].paragraphs[0].paragraph_format.space_after  = Pt(2)

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


# ── BUILD DOCUMENT ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── TITLE BLOCK ─────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(6)
r = p_title.add_run("The Defender Who Almost Nobody Noticed")
r.bold = True
r.font.size  = Pt(26)
r.font.color.rgb = HEAT_RED
r.font.name  = "Georgia"

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
r2 = p_sub.add_run(
    "How Hawk-Eye tracking data revealed the gap between looking good\n"
    "on defense and actually being good at it"
)
r2.italic = True
r2.font.size  = Pt(13)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
r2.font.name  = "Georgia"

p_byline = doc.add_paragraph()
p_byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_byline.paragraph_format.space_after = Pt(12)
r3 = p_byline.add_run("Judy Zhu  ·  May 2026  ·  judy.zhu6052@gmail.com")
r3.font.size  = Pt(10)
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r3.font.name  = "Calibri"

add_horizontal_rule(doc, "981F2A")
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── OPENING ─────────────────────────────────────────────────────────────────
body(doc, "Here's a stat that will break your brain a little.")
body(doc,
    "In the sample of Miami Heat games I analyzed, one defender allowed opponents to make only "
    "**20.8%** of their three-point attempts. The league average is around 36%. That's not just "
    "good — that's historically elite. The same defender's teammates were getting lit up for 44%, "
    "50%, even 55% on three-pointers.")
body(doc, "That defender? **Bam Adebayo.**")
body(doc,
    "Yes, the center. The guy whose job description, if you asked a casual fan, is \"not defend "
    "three-pointers.\" By every conventional stat — hustle metrics, crowd appreciation of "
    "closeouts, anything your eye naturally gravitates toward — Bam ranks last on his own team's "
    "perimeter defense leaderboard. He doesn't sprint across the floor to contest shots. He "
    "doesn't arrive at Hollywood speed.")
body(doc, "He's just… already there. Blocking the view. Arm up. Feet still.")
body(doc,
    "This story is about why that matters more than almost anything else, and what it took to "
    "actually see it.")

add_horizontal_rule(doc)

# ── SECTION 1 ───────────────────────────────────────────────────────────────
subheading(doc, "What Box Scores Can't Tell You")

body(doc,
    "The NBA tracks a stat called \"contest rate\" — how often a defender is within a certain "
    "distance of a shooter. It's useful the way a blurry photograph is useful. You get the "
    "general idea, but you're missing the things that actually distinguish a painting from a "
    "snapshot.")
body(doc,
    "Was the defender in front of the shooter, or chasing from behind? Were they arriving at "
    "the release, or had they been standing there for two seconds? Were their hands above the "
    "ball, or flailing somewhere around hip level? Did they jump at the pump fake?")
body(doc,
    "These questions can't be answered from a box score. They can't even be answered from video, "
    "not at scale. To answer them systematically across hundreds of possessions, you need "
    "something like Hawk-Eye.")
body(doc,
    "Hawk-Eye is an optical tracking system installed at NBA arenas — the same technology used "
    "in tennis and cricket to trace ball trajectories. In basketball, it captures the 3D "
    "position of every player and the ball at **60 frames per second**. Not just where they are "
    "standing, but where their wrists are. Where the ball is relative to the defender's hand. "
    "How fast the defender was moving in the half-second before the ball left the shooter's "
    "fingertips.")
body(doc,
    "I spent the past several weeks building a pipeline to extract that information from 15 "
    "Miami Heat games — about 1.1 gigabytes of tracking data per game — and using it to answer "
    "a question that sounds simple but turns out to be surprisingly hard:")

pull_quote(doc, "Who is actually defending three-point shots?")

add_horizontal_rule(doc)

# ── SECTION 2 ───────────────────────────────────────────────────────────────
subheading(doc, "Building a Metric From Scratch")

body(doc, "The first thing I had to decide was what \"actively contesting\" even means.")
body(doc,
    "This sounds obvious but isn't. If a defender is 9 feet away from a shooter at the release "
    "frame, are they contesting? What about 7 feet? What if they're 5 feet away but running "
    "directly perpendicular — beside the shooter, not in front?")
body(doc, "I settled on two filters grounded in basketball geometry. At the exact frame the ball leaves the shooter's hand:")

bullet(doc,
    "**Within 8 feet.** Beyond that distance, the research literature suggests the disruption "
    "effect is minimal — the shooter's mechanics aren't meaningfully altered.")
bullet(doc,
    "**Angle ≥ 90°.** The contest angle is the angle between the defender-to-shooter vector "
    "and the shooter-to-rim vector. At 90°, the defender is at the shooter's side. At 180°, "
    "they're standing directly between the shooter and the basket — a full face-guard. "
    "At 0°, they're behind the shooter.")

body(doc,
    "About **4% of apparent \"contests\"** in raw tracking data are defenders sprinting toward "
    "a shooter from behind — they show high speed and low distance, but they aren't blocking "
    "the shooter's sightline at release. They look good on a stat sheet. They feel like effort. "
    "They're not contesting the shot.")
body(doc,
    "After applying both filters: **310 actively contested three-point attempts** across 12 "
    "Heat defenders with at least 5 qualifying plays.")

add_horizontal_rule(doc)

# ── SECTION 3 ───────────────────────────────────────────────────────────────
subheading(doc, "Shot Contest Quality: Four Numbers, One Score")

body(doc,
    "To compare defenders, I built a composite metric called **Shot Contest Quality (SCQ)** — "
    "a 0-to-100 score computed at the shot release frame from four measurable dimensions:")

# SCQ formula table
data_table(doc,
    ["Component", "Weight", "What It Measures"],
    [
        ["Defender distance to shooter", "35%", "Proximity at release (closer = better)"],
        ["Closeout speed",               "30%", "Proxy for starting position — how far away the defender began"],
        ["Contest angle",                "20%", "How directly in front of the shooter (90° = side, 180° = face-guard)"],
        ["Hand height above ball",       "15%", "Defender's wrist relative to ball at release"],
    ],
    col_widths=[2.2, 0.7, 3.0],
)

body(doc,
    "A note on speed: this is not a measure of raw athleticism. A defender who closes out "
    "from 15 feet at full speed covers the ground — but they're still arriving having "
    "surrendered the advantage of being in position. High speed correlates with having "
    "started far away.")

body(doc, "Here's the full leaderboard, with each component visible:")

# SCQ rankings table
data_table(doc,
    ["Rank", "Defender", "Shots", "SCQ", "Speed pts", "Angle pts", "Hand pts"],
    [
        ["1",  "Dru Smith",             "16", "71.1", "28.9 ↑↑", "11.2 ↓↓", "14.9"],
        ["2",  "Kasparas Jakucionis",   "13", "69.6", "26.5 ↑",  "13.1",    "14.6"],
        ["3",  "Davion Mitchell",       "48", "69.5", "25.3 ↑",  "14.3",    "13.9"],
        ["4",  "Tyler Herro",           "32", "68.8", "25.7 ↑",  "13.7",    "13.3"],
        ["4",  "Kel'el Ware",           "10", "68.8", "24.1",    "16.3 ↑↑", "15.0 ↑↑"],
        ["6",  "Jaime Jaquez Jr.",      "40", "67.1", "23.6",    "14.2",    "14.5"],
        ["7",  "Pelle Larsson",         "39", "67.0", "23.7",    "14.1",    "14.1"],
        ["8",  "Simone Fontecchio",     "10", "66.8", "27.6 ↑",  "12.4 ↓",  "14.4"],
        ["9",  "Norman Powell",         "15", "66.0", "22.4",    "15.2 ↑",  "14.2"],
        ["10", "Nikola Jovic",          "10", "65.1", "21.5",    "14.3",    "15.0 ↑↑"],
        ["11", "Andrew Wiggins",        "52", "64.6", "20.5 ↓",  "16.1 ↑↑", "14.6"],
        ["12", "Bam Adebayo",           "25", "62.3", "17.6 ↓↓", "15.6 ↑",  "15.0 ↑↑"],
    ],
    col_widths=[0.45, 1.7, 0.5, 0.55, 0.85, 0.85, 0.85],
)

body(doc, "*Pool average SCQ ≈ 67. ↑↑ = notably above average; ↓↓ = notably below average.*", italic=True)
doc.add_paragraph()

body(doc,
    "Two things pop immediately. **Dru Smith** is the team's best contest player — by speed. "
    "His speed score (28.9) is the highest on the team by a wide margin. He is phenomenal at "
    "sprint-closing out from distance. But his **angle score (11.2) is the worst on the team**. "
    "He arrives fast, but he arrives from the side.")
body(doc,
    "**Bam Adebayo** is last in overall SCQ (62.3), entirely due to his speed score (17.6). "
    "He almost never closes out — because he's almost never out of position. His angle (15.6) "
    "and hand (15.0) scores are both above average. He is correctly positioned and physically "
    "dominant when his man catches and shoots.")

add_horizontal_rule(doc)

# ── SECTION 4 ───────────────────────────────────────────────────────────────
subheading(doc, "When the Leaderboard Lies")

body(doc,
    "SCQ tells you who contests best. But contesting well and *actually suppressing makes* are "
    "different things — and the gap between them is one of the most important insights this "
    "data produces.")
body(doc,
    "To measure actual defensive contribution, I built a **lift framework**. Before crediting "
    "a defender for a miss, account for what kind of shooter and shot you're talking about. "
    "A 34% shooter taking a step-back pull-up from 27 feet misses a lot — but that might "
    "have nothing to do with the defender.")
body(doc,
    "I trained a ridge logistic regression model on the shooter side only: the shooter's "
    "season three-point percentage, distance to the rim, shot height, arc, and shot clock "
    "remaining. This gives an *expected* make probability for every shot, with no defensive "
    "information. **Lift = actual make rate minus expected make rate.** Negative lift is "
    "genuine shot suppression.")

# Lift table
data_table(doc,
    ["Defender", "Shots", "Actual %", "Expected %", "Lift"],
    [
        ["Bam Adebayo",          "24", "20.8%", "41.3%", "−20.5 pp ★★★"],
        ["Tyler Herro",          "32", "34.4%", "45.7%", "−11.4 pp ★★★"],
        ["Kasparas Jakucionis",  "13", "30.8%", "40.8%", "−10.0 pp ★★"],
        ["Norman Powell",        "15", "40.0%", "46.9%", "−6.9 pp ★★"],
        ["Kel'el Ware",          "10", "40.0%", "45.6%", "−5.6 pp ★"],
        ["Simone Fontecchio",    "9",  "44.4%", "43.1%", "+1.3 pp  ∼"],
        ["Andrew Wiggins",       "50", "44.0%", "41.9%", "+2.1 pp  ∼"],
        ["Dru Smith",            "16", "50.0%", "43.1%", "+6.9 pp  −"],
        ["Davion Mitchell",      "48", "52.1%", "45.1%", "+7.0 pp  −"],
        ["Pelle Larsson",        "38", "52.6%", "44.6%", "+8.0 pp  −"],
        ["Jaime Jaquez Jr.",     "38", "55.3%", "43.4%", "+11.9 pp −"],
    ],
    col_widths=[1.85, 0.55, 0.85, 0.95, 1.15],
)

body(doc, "*pp = percentage points. ★★★ strong signal; ★ directional only (n < 20); ∼ neutral; − above expected.*", italic=True)
doc.add_paragraph()

body(doc,
    "The Spearman correlation between SCQ rank and lift rank is **+0.34** — weakly positive, "
    "meaning higher-contest-quality defenders are slightly associated with *worse* outcomes. "
    "This is a paradox, but not a mystery.")

pull_quote(doc,
    "When a team's best wing defenders are on the floor, the coaching staff isn't randomly "
    "distributing coverage. They're being sent to the most dangerous shooters in the most "
    "dangerous situations. High-SCQ defenders are in hard situations because they're trusted "
    "to be there.")

body(doc,
    "At the individual shot level, harder contests do suppress makes — the shot-level "
    "Spearman is −0.078, meaning the relationship exists in the right direction. But at the "
    "player aggregate level, the defenders making the most hard contests are doing so partly "
    "because they're assigned to the opponents who create the most hard situations.")
body(doc,
    "The lift framework resolves this by design. By accounting for shooter quality and shot "
    "difficulty before measuring outcomes, it separates \"you contested a hard situation well\" "
    "from \"you were put in hard situations.\"")

add_horizontal_rule(doc)

# ── SECTION 5 ───────────────────────────────────────────────────────────────
subheading(doc, "The Physical Model: Size Matters More Than Speed")

body(doc,
    "The most technically interesting part of this analysis is what happens when you add "
    "physical dimensions to the prediction model.")
body(doc,
    "I matched each player to their NBA combine measurements and engineered three additional "
    "features from the tracking data:")

bullet(doc,
    "**Effective contest height**: How many inches above (or below) the ball is the "
    "defender's highest wrist at the release frame? This directly captures whether the "
    "hand is obstructing the shooter's release window.")
bullet(doc,
    "**Height difference**: Defender's height minus shooter's height, in inches.")
bullet(doc,
    "**Wingspan vs. shooter height**: Defender's wingspan minus shooter's height — "
    "a reach advantage measure independent of the defender's own proportions.")

body(doc, "When I added these to the full model alongside contest mechanics, here's what emerged:")

data_table(doc,
    ["Feature", "Odds Ratio per +1 SD", "Interpretation"],
    [
        ["Effective contest height", "0.882 (strongest)", "Each SD of hand-over-ball → −12% make odds"],
        ["Height difference",        "0.885",             "Raw height advantage nearly as powerful"],
        ["Wingspan advantage",       "0.908",             "Reach over shooter, independent of height"],
        ["SCQ composite",            "0.912",             "Combined contest quality"],
        ["Defender jump (pre-release)", "1.138 ↑",        "Larger jumps → MORE makes (pump-fake artifact)"],
    ],
    col_widths=[2.0, 1.5, 2.9],
)

doc.add_paragraph()

body(doc,
    "The **defender jump** result deserves explanation. Shouldn't jumping help? The issue "
    "is timing. The 250ms pre-release window catches defenders who already left the floor — "
    "responding to a pump fake, misjudging the release point, or going too early. "
    "The shots that go in most often in this data are the ones where the defender jumped "
    "and the shooter waited them out, then released into a completely uncontested window.")

pull_quote(doc,
    "The best contests in this dataset are disciplined, grounded hand-extensions. "
    "Not explosive vertical efforts.")

body(doc,
    "This has real implications for player development. You cannot coach a player to be "
    "taller. You can coach them not to leave the floor.")

add_horizontal_rule(doc)

# ── SECTION 6 ───────────────────────────────────────────────────────────────
subheading(doc, "Sensitivity Analysis: The Weighting Problem")

body(doc,
    "Any time you build a composite metric, a reasonable critic will ask: *what if you chose "
    "different weights?* I tested five alternative weighting schemes and measured how the "
    "rankings changed:")

data_table(doc,
    ["Weighting Scheme", "dist", "speed", "angle", "hand", "Spearman ρ vs. Original"],
    [
        ["Original (baseline)",         "0.35", "0.30", "0.20", "0.15", "1.000"],
        ["Equal weights",               "0.25", "0.25", "0.25", "0.25", "+0.916"],
        ["Distance-heavy",              "0.55", "0.20", "0.15", "0.10", "+0.937"],
        ["Technique-heavy (angle+hand)","0.20", "0.15", "0.35", "0.30", "−0.126"],
        ["No speed (stationed focus)",  "0.45", "0.00", "0.30", "0.25", "−0.462"],
    ],
    col_widths=[2.1, 0.45, 0.55, 0.55, 0.55, 1.6],
)
doc.add_paragraph()

body(doc,
    "Those last two are not typos. When you weight angle and hand heavily — or remove speed "
    "entirely — the rankings don't shift slightly. **They invert.** Dru Smith goes from #1 "
    "to #10. Bam Adebayo goes from #12 to #6. Andrew Wiggins goes from #11 to #2.")
body(doc, "This isn't metric fragility. It's the metric exposing a genuine philosophical split:")

data_table(doc,
    ["Defensive Profile", "Examples", "Strength", "Weakness"],
    [
        ["Closer",    "Dru Smith, Fontecchio, Jakucionis", "High speed, good distance",       "Angle: arriving from the side"],
        ["Stationed", "Kel'el Ware, Adebayo, Wiggins",    "Best angle, best hand position",  "Low speed (already there)"],
    ],
    col_widths=[1.2, 2.2, 1.8, 1.8],
)
doc.add_paragraph()

body(doc,
    "*Who closes out hardest?* favors speed. *Who stands in the most disruptive position?* "
    "favors angle and hand. A complete evaluation uses both — and recognizes they're "
    "measuring different defensive styles, not the same one on a single continuum.")

add_horizontal_rule(doc)

# ── SECTION 7 ───────────────────────────────────────────────────────────────
subheading(doc, "What This Means for the Front Office")

body(doc,
    "If you're a team evaluating a free agent wing defender, what do you look at? If you "
    "look at opponent three-point percentage, you're absorbing the noise of who they guarded. "
    "If you look at contest rate, you're measuring activity, not quality. If you look at "
    "highlight reels, you're seeing the plays someone chose to put in a video.")
body(doc, "The framework here gives you three things highlights don't:")

bullet(doc,
    "**Contest quality, not contest quantity.** Smith and Jakucionis generate high-quality "
    "contests, not just high-frequency ones. Their SCQ scores are robust across almost every "
    "weighting assumption that includes speed.")
bullet(doc,
    "**Outcome controlling for assignment difficulty.** Tyler Herro's −11.4 pp lift is "
    "genuinely impressive. He's being deployed in high-pressure situations and converting "
    "hard closeouts into misses. Raw make rate doesn't tell you that.")
bullet(doc,
    "**A physical scouting signal that isn't in the box score.** Effective contest height — "
    "how high the hand is above the ball at the moment it leaves the shooter's hand — is "
    "more predictive than any speed or athleticism measure in this dataset. Combined with "
    "the discipline to not leave the floor, it defines the highest-value contest profile.")

add_horizontal_rule(doc)

# ── SECTION 8 ───────────────────────────────────────────────────────────────
subheading(doc, "A Note on What 15 Games Can and Can't Tell You")

body(doc,
    "I want to be honest about the limits of this analysis, because that's actually part of "
    "what makes it useful.")
body(doc,
    "310 shots across 12 defenders is enough to identify patterns, but not enough to be "
    "confident in individual player conclusions. Bam Adebayo's −20.5 pp lift is the most "
    "dramatic number in the dataset, but 24 shots is a small sample. The direction is "
    "meaningful; the magnitude should be treated as directional.")
body(doc,
    "I would not walk into a front office meeting and say \"Jaquez is a bad three-point "
    "defender.\" I would say: \"In 15 games, he allowed 12 more makes per hundred contests "
    "than shooter skill and shot difficulty would predict, and here's what the tracking data "
    "shows about his contest mechanics in those situations.\"")
body(doc, "That's a different conversation. And it's the kind of conversation this data makes possible.")

pull_quote(doc,
    "A full season of Hawk-Eye data — 82 games, 2,000+ contested attempts — would turn "
    "directional patterns into something you could act on with confidence. The framework "
    "is built. The question is just how much data you pour into it.")

add_horizontal_rule(doc)

# ── CLOSING ─────────────────────────────────────────────────────────────────
subheading(doc, "The Quiet Defenders")

body(doc,
    "The single most interesting basketball insight from this analysis is also the simplest: "
    "the Heat's best three-point defense in this sample comes from players who look boring "
    "doing it.")
body(doc,
    "Bam Adebayo doesn't sprint across the floor. He doesn't make the SportsCenter closeout. "
    "He walks into position, gets his hand up, and doesn't jump when the shooter pump fakes. "
    "The result is a 20-percentage-point suppression — the kind of number that, maintained "
    "over a full season, would be among the best defensive performances in the league.")
body(doc,
    "Tyler Herro, a player most people think of as a scorer, shows up at −11 pp. His "
    "technical contest quality is middle-of-the-pack, but he's positioned correctly and his "
    "hand gets above the ball. He's not the story anyone tells about Heat defense.")
body(doc,
    "The noisiest, most athletic, most visible defenders — the sprint closers, the "
    "spectacular recovery artists — are the ones the camera loves. They're also the ones "
    "associated with more makes than expected in this data.")
body(doc,
    "There's a version of this that's unfair to closers. They are doing genuinely hard things "
    "in genuinely hard situations. The matchup selection problem is real. But there is also a "
    "version of this that is just basketball truth.")

pull_quote(doc,
    "Being in position is better than recovering to position. A hand above the ball matters "
    "more than the speed of the feet that got it there.")

body(doc,
    "The data just takes a while to see what the smart coaches probably already know.")

add_horizontal_rule(doc)

# ── FOOTER NOTE ─────────────────────────────────────────────────────────────
p_footer = doc.add_paragraph()
p_footer.paragraph_format.space_before = Pt(8)
_add_inline(p_footer,
    "*Judy Zhu is a sports analytics researcher. This analysis was conducted using Hawk-Eye "
    "optical tracking data from 15 Miami Heat games, 2025–26 NBA regular season. "
    "All modeling implemented from scratch in Python. If you work in a front office and "
    "found this interesting: judy.zhu6052@gmail.com*")
for run in p_footer.runs:
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(OUT)
print(f"Saved → {OUT}")
